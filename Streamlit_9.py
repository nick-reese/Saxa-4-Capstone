import pandas as pd
import numpy as np
import streamlit as st
from docx import Document
import pdfplumber
import csv
import io
import os
import json
from textwrap import dedent
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.corpus import stopwords
import warnings
warnings.filterwarnings("ignore")
from dotenv import load_dotenv
import openai
from openai import OpenAI
from sklearn.metrics.pairwise import cosine_similarity 
import spacy
# import spacy_streamlit
# models = ["en_core_web_sm", "en_core_web_md"]
nlp = spacy.load("en_core_web_lg")
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
import re


#####################################################################

st.image('georgetown_image.jpeg', use_container_width=True)



st.markdown(
    """
    <style>
    .main-caption {
        font-size: 24px;
        font-weight: bold;
        text-align: center;
        color: #333;
    }
    .name-list {
        font-size: 18px;
        text-align: center;
        color: #555;
    }
    </style>
    <p class="main-caption">Saxa - 4</p>
    <p class="name-list">Nicholas Reese</p>
    <p class="name-list">Ashlyn Bellardine</p>
    <p class="name-list">Osama Bin Habibi</p>
    <p class="name-list">Dezmond Richardson</p>
    <p class="name-list">Genesis Roberto</p>
    
    """,
    unsafe_allow_html=True
)



st.markdown('---')
#####################################################################
csv_file_path = 'spacy_redacted_documents_with_id_and_category.csv'
data = pd.read_csv(csv_file_path)
json_file_path = 'spacy_redacted_documents_with_id_and_category.json'

data.to_json(json_file_path, orient = 'records', lines = True)

resumes = data

# This below is used for reading in anytype of resume.
st.markdown('# Recommendation System')
st.markdown('---')
#####################################################################

st.write('## Section 1 - Loading Resume')
st.markdown('### Upload your Resume')

uploaded_file = st.file_uploader('Drag & and Drop your resume. We can analyze word, pdf, txt or csv',
                                 type = ['docx', 'txt', 'pdf', 'csv'])


# Word
if uploaded_file is not None:
    resume_df = None
    def read_docx(file):
        doc = Document(file)
        data = [para.text for para in doc.paragraphs]
        return pd.DataFrame(data, columns = ['Redacted Text'])
# Text
    def read_txt(file):
        data = file.read().decode("utf-8").splitlines()
        return pd.DataFrame(data, columns = ['Redacted Text'])
# Pdf
    def read_pdf(file):
        with pdfplumber.open(file) as pdf:
            data = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    data.extend(text.splitlines())
                return pd.DataFrame(data,  columns= ['Redacted Text'])
# CSV

    def read_csv(file):
        df = pd.read_csv(file)
        return df
    
    if uploaded_file.name.endswith('.docx'):
        resume_df = read_docx(uploaded_file)
    elif uploaded_file.name.endswith('.txt'):
        resume_df = read_txt(uploaded_file)
    elif uploaded_file.name.endswith('.pdf'):
        resume_df = read_pdf(uploaded_file)
    elif uploaded_file.name.endswith('.csv'):
        resume_df = read_csv(uploaded_file)
    else:
        st.error("Unsupported file type!")


st.markdown('### Non Processed Resume')
st.dataframe(resume_df)

#######################################################################

st.markdown('---')
st.markdown('## Section 2 - LLM')

st.markdown('### Recommender Output')
st.markdown('#### Below the Recommendation will be using a default prompt selected by the team')

st.markdown('---')
st.write("""
       You are a helpful recommender tool. You will be provided a resume where the individual would like to have a different but similar
        job recommended to them based off their resume. Your goal will be to provide the five most similar jobs to what is on their resume that they provided.
        For each recommendation, please provide an explanation as to why you chose the job you did for the individual.
        """)

st.markdown('---')
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

#######################################################################
# Adding Stopwords and Vectorizing 

sw = stopwords.words("english")
sw.extend(['[Redacted]'])

vec = TfidfVectorizer(stop_words = sw)

################################

# Calculating Distance 


def calc_sim(resume_df, resumes):
    
    new_resume = str(resume_df)
    
    resumes['Redacted Text'] = resumes['Redacted Text'].astype(str)
    
    all_resumes = resumes['Redacted Text'].tolist() + [new_resume]
    
    vec = TfidfVectorizer(stop_words = sw)
    tfidf_matrix = vec.fit_transform(all_resumes)
    
    similarity_matrix = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1])
    
    similar_indices = similarity_matrix.argsort()[0][-5:][::-1]
    
    return similar_indices

##################################################
# Adding in the function to get recs

resume_rec_prompt = '''
        You are a helpful recommender tool. You will be provided a resume where the individual would like to have a different but similar
        job recommended to them based off their resume. Your goal will be to provide the five most similar jobs to what is on their resume that they provided.
        For each recommendation, please provide an explanation as to why you chose the job you did for the individual.

        For example this is the structure that the output should be as follows:
        
        Based on the resume provided, here are five job recommendations for the individual:

        1) Data Scientist: With experience in data analysis, modeling, and visualization, as well as proficiency in R, Python, and statistical software, 
        the individual would be well-suited for a role as a Data Scientist. They have the skills necessary to work with diverse sources 
        of data and develop predictive models.
        '''
new_resume = resume_df

def get_rec_roles(new_resume, resumes):
    
    #new_resume = resume_df
    similar_indices = calc_sim(new_resume, resumes)
    
    resumes_texts = resumes['Redacted Text'].tolist()
    
    selected_resumes = [resumes_texts[i] for i in similar_indices]
    
    prompt = resume_rec_prompt + '\n\n' + '\n'.join(selected_resumes)
    
    try:
        response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": prompt}
        ],
        temperature= 0.3
    )
        return response.choices[0].message.content
    except Exception as e:
        print(f' an error occurred at {e}')
        return None

result = get_rec_roles(new_resume, resumes)

st.markdown(result)

################################

## Adding the Default Prompt
default_prompt = '''
You are a helpful recommender tool. You will be provided a resume where the individual would like to have a different but similar
job recommended to them based off their resume. Your goal will be to provide five similar jobs to what is on their resume that they provided.
For each reommendation, please provide an explanation as to why you chose the job you did for the individual.
'''
#####################################

def calc_sim(resume_df, resumes):
    
    new_resume = str(resume_df)
    
    resumes['Redacted Text'] = resumes['Redacted Text'].astype(str)
    
    all_resumes = resumes['Redacted Text'].tolist() + [new_resume]
    
    vec = TfidfVectorizer(stop_words = sw)
    tfidf_matrix = vec.fit_transform(all_resumes)
    
    similarity_matrix = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1])
    
    similar_indices = similarity_matrix.argsort()[0][-5:][::-1]
    
    return similar_indices



def get_rec_roles2(new_resume, resumes, user_prompt):
    new_resume = resume_df
    similar_indices = calc_sim(new_resume, resumes)
    resumes_texts = resumes['Redacted Text'].tolist()
    selected_resumes = [resumes_texts[i] for i in similar_indices]

    prompt = user_prompt + '\n\n' + '\n'.join(selected_resumes)
    try:
        response = client.chat.completions.create(
            model = 'gpt-3.5-turbo',
            messages= [
                {'role': 'user',
                 'content': prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f'An error occured: {e}')
        return None

st.markdown('---')


def main():
    st.markdown("### Choose your own prompt")

    user_prompt = st.text_area('Below you can ask your own prompt. If you have a specific industry in mind, you can see what skills might help you acquire a role', 
                               value= default_prompt)
    
    if st.button('Get Recommendations'):
        if new_resume is not None:
            final_prompt = user_prompt.strip() if user_prompt.strip() else default_prompt
            
            recommendations = get_rec_roles2(new_resume, resumes, final_prompt)
            if recommendations:
                st.write('### Recommendations:')
                st.write(recommendations)
            else:
                st.write('No Recs were generated')
        else: 
            st.write('Please enter your resume')
if __name__ == '__main__':
    main()    

st.markdown('---')
#######################################################################

st.markdown('## Section 3 - PII')

st.write("""
         In issue that the team wanted to tackle with their Capstone was the redaction of Personally Identifiable Information or PII. There are many reasons why the team may want to redact PII,
         but some that were the most prevelent were bias, privacy and accountability. 

         The team wanted to test whether names of inidividuals, names of schools, names of former employeers, locations and other identifiable information produced a different output for a resume that had 
         PII redacted and a resume without PII redacted. 

         The same can be said about privacy and accountability, the team wanted to highlight the importance of keeping your information private. 
         """)

st.markdown('---')

nlp = spacy.load("en_core_web_lg")

#new_resume_text = resume_df

def ner_pipeline(text):
        doc = nlp(text)
        return [{'word': ent.text, 'entity': ent.label_, 'score': ent.kb_id_ if ent.kb_id_ else 1.0} for ent in doc.ents]
        #return entities


def redaction(new_resume_text):
    st.markdown('### Resume Redaction')
    
    entities = ner_pipeline(new_resume_text)
    thresholds = {
        'PERSON': 0.75,
        'ORG': 0.99,
        'GPE': 0.99, 
        }

    redacted_text = new_resume_text

    for entity in sorted(entities, key= lambda x: len(x['word']), reverse= True):
        score = float(entity.get('score', 1.0))
        if score >= thresholds.get(entity['entity'], 0):
            redacted_text = re.sub(re.escape(entity['word']), '[Redacted]', 
                                   redacted_text)

    #entities_sorted = sorted(entities, key= lambda x:len(x['word']), reverse= True)
    #for entity in entities_sorted:
     #   entity_text = entity['word']
      #  entity_label = entity['entity']
       # score = entity.get('score', 1.0)

    
    redacted_text = re.sub(r'\S+@\S+', '[Redacted Email]', redacted_text)
    redacted_text = re.sub(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', '[Redacted Phone]', redacted_text)
    redacted_text = re.sub(r'\b(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9._-]+\b', '[Redacted Website]', redacted_text)
    
    return pd.DataFrame({
        'Redacted Text': [redacted_text]
    })
    #return spacy_redacted_resume


def extract_text_from_pdf(file):
    """Extracts text from a PDF file using pdfplumber."""
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file):
    """Extracts text from a Word document."""
    text = ""
    doc = Document(file)
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


#st.title("Resume Redaction")
#st.title("Resume Redaction Tool")


if uploaded_file is not None:
    if uploaded_file.type == "application/pdf":
        resume_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        resume_text = extract_text_from_docx(uploaded_file)
    
    
    st.markdown("### Extracted Resume Text")
    st.text_area("Resume Text", resume_text, height=300)

    redacted_df = redaction(resume_text)
    
    
    #st.markdown("### Redacted Resume")
    st.dataframe(redacted_df)  
else:
    st.write("Please upload a PDF or Word document to see the extracted text and redacted version.")




#######################################################################
st.write('''
         Now take a look and see what was redacted from your resume. Here we can see that our NER model
         is redacting information like your location, name, email, as well as the name of the schools or schools
         you may have gone to. 
         ''')
st.markdown('---')

#######################################################################

st.markdown('### Section 4 - Recommendation With Redaction')

st.markdown('---')


def get_rec_roles_redaction(redacted_df, resumes):
    
    #new_resume = resume_df
    similar_indices = calc_sim(redacted_df, resumes)
    
    resumes_texts = resumes['Redacted Text'].tolist()
    
    selected_resumes = [resumes_texts[i] for i in similar_indices]
    
    prompt = resume_rec_prompt + '\n\n' + '\n'.join(selected_resumes)
    
    try:
        response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": prompt}
        ],
        temperature= 0.5
    )
        return response.choices[0].message.content
    except Exception as e:
        print(f' an error occurred at {e}')
        return None

result2 = get_rec_roles_redaction(redacted_df, resumes)


st.markdown(result2)
st.markdown('---')

##########################################################################

st.markdown('### Section 5 - Redaction Recommendation with Prompt Enhancement')
st.write('''
         What we are seeing above is that the recommendations may have some "hallunications" based off how the similarity distance is calculated.
         Specifically, the team believes this effect may be coming from the LLM's lack of knowledge that the resumes have redacted
         information. As such the prompt below updates provides the LLM with additional knowledge. 
         ''')


st.write('''
         You are a helpful recommender tool. You will be provided a resume where the individual would like to have a different but similar
        job recommended to them based off their resume. Your goal will be to provide the five most similar jobs to what is on their resume that they provided.
        For each recommendation, please provide an explanation as to why you chose the job you did for the individual. The resume will have all 
        PII (personal Identifiable information) redacted and this will be shown as "[Redacted]" please keep this in mind when recommending. 

        ''')

st.markdown('---')


redacted_prompt = '''
You are a helpful recommender tool. You will be provided a resume where the individual would like to have a different but similar
        job recommended to them based off their resume. Your goal will be to provide the five most similar jobs to what is on their resume that they provided.
        For each recommendation, please provide an explanation as to why you chose the job you did for the individual.The resume will have all 
        PII (personal Identifiable information) redacted and this will be shown as "[Redacted]" please keep this in mind when recommending. 


        For example this is the structure that the ouput should be as follows:
        
        Based on the resume provided, here are five job recommendations for the individual:

        1) Data Scientist: With experience in data analysis, modeling, and visualization, as well as proficiency in R, Python, and statistical software, 
        the individual would be well-suited for a role as a Data Scientist. They have the skills necessary to work with diverse sources 
        of data and develop predictive models.

'''

def get_rec_roles_redaction_with_redacted_prompt(redacted_df, resumes):
    
    #new_resume = resume_df
    similar_indices = calc_sim(redacted_df, resumes)
    
    resumes_texts = resumes['Redacted Text'].tolist()
    
    selected_resumes = [resumes_texts[i] for i in similar_indices]
    
    prompt = resume_rec_prompt + '\n\n' + '\n'.join(selected_resumes)
    
    try:
        response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": redacted_prompt}
        ],
        temperature= 0.5
    )
        return response.choices[0].message.content
    except Exception as e:
        print(f' an error occurred at {e}')
        return None

result3 = get_rec_roles_redaction_with_redacted_prompt(redacted_df, resumes)
st.markdown(result3)
#########################################

st.markdown('---')

st.markdown('### Section 6 - Data Privacy Concerns')

st.write('''

        When it comes to a consumers choice, the team wanted to ensure that the individual had the decision to allow
        the use of their information, regardless of the redaction. This attribute of privacy transparency guarantees
        ethical use of their data.  
        ''')

st.markdown('---')

json_file_path = '/Users/nicholasreese/Desktop/Georgetown/Capstone/capstone_github/Capstone_Introduction/Saxa-4-Capstone/redacted_resumes_output.json'

def add_resume_to_database(redacted_df, json_file_path):
    if os.path.exists(json_file_path):
        existing_data = pd.read_json(json_file_path, lines= True)
    else:
        existing_data = pd.DataFrame()
    
    combined_data = pd.concat([existing_data, redacted_df], ignore_index= True).drop_duplicates()
    if len(combined_data) == len(existing_data):
        return False
    else:
        combined_data.to_json(json_file_path, orient= 'records', lines=True)
        return True

def main():
    st.subheader('To better our recomendations, could we add your resume?')
    if st.button('Yes'):
        added = add_resume_to_database(redacted_df, json_file_path)
        if added: 
            st.success('The redacted resume has been added. Thank you!')
        else:
            st.warning('This resume has already been added. Thank you for your ongoing support!')
    elif st.button('No'):
        st.info('This resume has not been added, thank you for using our product')

if __name__ == '__main__':
    main()

